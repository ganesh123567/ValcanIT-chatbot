from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document


TOKEN_RE = re.compile(r"[a-z0-9]+")
DEFAULT_DOCUMENTS_DIR = Path("data/documents")
DEFAULT_KNOWLEDGE_PATH = Path("data/valcanit_knowledge.json")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Ingest Word documents into the ValcanIT RAG knowledge base.")
    parser.add_argument("--documents-dir", default=str(DEFAULT_DOCUMENTS_DIR), help="Folder containing .docx files.")
    parser.add_argument("--knowledge-path", default=str(DEFAULT_KNOWLEDGE_PATH), help="Knowledge JSON to update.")
    parser.add_argument("--chunk-size", type=int, default=900, help="Approximate max characters per chunk.")
    return parser.parse_args()


def docx_text(path: Path) -> str:
    document = Document(path)
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def chunk_text(text: str, chunk_size: int) -> list[str]:
    paragraphs = [paragraph.strip() for paragraph in text.splitlines() if paragraph.strip()]
    chunks: list[str] = []
    current = ""

    for paragraph in paragraphs:
        candidate = f"{current}\n{paragraph}".strip()
        if current and len(candidate) > chunk_size:
            chunks.append(current)
            current = paragraph
        else:
            current = candidate

    if current:
        chunks.append(current)

    return chunks


def keywords_for(source_name: str, content: str) -> list[str]:
    tokens = TOKEN_RE.findall(f"{source_name} {content}".lower())
    seen: set[str] = set()
    keywords: list[str] = []
    for token in tokens:
        if len(token) < 3 or token in seen:
            continue
        seen.add(token)
        keywords.append(token)
    return keywords[:40]


def load_knowledge(path: Path) -> list[dict]:
    if not path.exists():
        return []
    with path.open(encoding="utf-8") as file:
        return json.load(file)


def main() -> None:
    args = parse_args()
    documents_dir = Path(args.documents_dir)
    knowledge_path = Path(args.knowledge_path)

    knowledge = load_knowledge(knowledge_path)
    documents = sorted(documents_dir.glob("*.docx"))
    if not documents:
        print(f"No .docx files found in {documents_dir}")
        return

    source_files = {str(path) for path in documents}
    knowledge = [item for item in knowledge if item.get("source_file") not in source_files]

    added = 0
    for document_path in documents:
        text = docx_text(document_path)
        chunks = chunk_text(text, args.chunk_size)
        for index, chunk in enumerate(chunks, start=1):
            title = document_path.stem.replace("_", " ").replace("-", " ").title()
            if len(chunks) > 1:
                title = f"{title} - Part {index}"
            knowledge.append(
                {
                    "title": title,
                    "content": chunk,
                    "keywords": keywords_for(document_path.stem, chunk),
                    "source_file": str(document_path),
                }
            )
            added += 1

    with knowledge_path.open("w", encoding="utf-8") as file:
        json.dump(knowledge, file, indent=2)
        file.write("\n")

    print(f"Ingested {len(documents)} Word document(s) into {knowledge_path}.")
    print(f"Added {added} RAG chunk(s).")
    print("Restart FastAPI so the in-memory RAG cache reloads.")


if __name__ == "__main__":
    main()
